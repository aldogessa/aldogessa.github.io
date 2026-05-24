"""
==========================================================================
AGis - Aldo Gessa
SCRIPT – QGIS 3.40.11 (versione GeoPackage)
Crea il CDU dalla selezione dei record nella tabella precalcolata
delle destinazioni urbanistiche.
La struttura del layer deve essere quella del workflow
Maggio 2026
==========================================================================
"""

from qgis.PyQt.QtWidgets import QFileDialog, QMessageBox
from qgis.core import QgsProject
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import os
import platform
import re
from qgis.utils import iface

def parse_mappale(value):
    v = str(value).strip()

    # Caso 1: solo numeri → "12"
    if v.isdigit():
        return (0, int(v), "")

    # Caso 2: numeri + lettere → "12A", "3B"
    match = re.match(r"^(\d+)([A-Za-z]+)$", v)
    if match:
        num = int(match.group(1))
        suf = match.group(2)
        return (1, num, suf)

    # Caso 3: solo lettere → "A", "B"
    if v.isalpha():
        return (2, 0, v)

    # Caso 4: casi strani → ordina come stringa pura
    return (3, 0, v)

def export_cdu_from_destination_layer():
    try:
        # Recupera il layer dell'azione
        layer_id = "[% @layer_id %]"
        layer = QgsProject.instance().mapLayer(layer_id)

        if not layer:
            iface.messageBar().pushMessage("Errore", "Impossibile identificare il layer dell'azione.", level=3)
            return

        # Controllo selezione
        selected = layer.selectedFeatures()
        if not selected:
            iface.messageBar().pushMessage("Avviso", "Nessun record selezionato nel layer.", level=2)
            return

        # Scelta ORDINARIO / COMPLETO
        msg = QMessageBox()
        msg.setWindowTitle("Tipo di certificato")
        msg.setText("Desideri il certificato ORDINARIO (filtrato) o COMPLETO (tutti i dati)?")

        btn_ordinario = msg.addButton("ORDINARIO", QMessageBox.ActionRole)
        btn_completo = msg.addButton("COMPLETO", QMessageBox.ActionRole)
        btn_annulla = msg.addButton("ANNULLA", QMessageBox.RejectRole)

        msg.exec_()

        if msg.clickedButton() == btn_annulla:
            iface.messageBar().pushMessage("Informazione", "Operazione annullata.", level=1)
            return

        applica_filtro = (msg.clickedButton() == btn_ordinario)

        # Finestra salvataggio
        output_file_path, _ = QFileDialog.getSaveFileName(
            None,
            "Scegli dove salvare il CDU",
            "",
            "Documenti Word (*.docx)"
        )
        if not output_file_path:
            iface.messageBar().pushMessage("Informazione", "Operazione annullata.", level=1)
            return

        # Template Word
        template_path = os.path.join(QgsProject.instance().homePath(), "Immagini", "00_CDU_Schema.docx")
        if not os.path.exists(template_path):
            iface.messageBar().pushMessage("Errore", f"Template non trovato: {template_path}", level=3)
            return

        doc = Document(template_path)

        # Tabella Word
        if not doc.tables:
            iface.messageBar().pushMessage("Errore", "Nessuna tabella trovata nel template!", level=3)
            return

        table = doc.tables[0]

        # Intestazioni
        headers = ["Fg.", "All.", "Map.", "Tema", "Zona", "Dettaglio", "Norme Specifiche", "Q.tà* %"]

        # Allineamenti intestazione e dati
        allineamenti = [
            WD_ALIGN_PARAGRAPH.CENTER,   # Fg.
            WD_ALIGN_PARAGRAPH.CENTER,   # All.
            WD_ALIGN_PARAGRAPH.CENTER,   # Map.
            WD_ALIGN_PARAGRAPH.LEFT,     # Tema
            WD_ALIGN_PARAGRAPH.CENTER,   # Zona
            WD_ALIGN_PARAGRAPH.LEFT,     # Dettaglio
            WD_ALIGN_PARAGRAPH.LEFT,     # Norme
            WD_ALIGN_PARAGRAPH.CENTER    # Q.tà %
        ]

        # Funzione per aggiungere intestazione
        def aggiungi_intestazione():
            header_row = table.add_row()
            for i, header in enumerate(headers):
                cell = header_row.cells[i]
                cell.text = header

                for paragraph in cell.paragraphs:
                    run = paragraph.runs[0]
                    run.bold = True
                    run.font.name = "Calibri Light"
                    run.font.size = Pt(10)
                    paragraph.alignment = allineamenti[i]

                # Rimuovi eventuale sfondo
                tcPr = cell._element.get_or_add_tcPr()
                for shd in tcPr.findall(".//w:shd", namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}):
                    tcPr.remove(shd)

                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        # Codici ammessi per ORDINARIO
        codici_ammessi = ("T01", "T02", "T03")

        # Ordina i record selezionati
        selected_sorted = sorted(
            selected,
            key=lambda f: (
                int(f["FOGLIO"]),
                parse_mappale(f["MAPPALE"]),
                f["TEMA"],
                f["ZONA"],
                f["PERCENT_V"]
            )
        )

        ultimo_gruppo = None

        # Popolamento tabella
        for f in selected_sorted:

            # Filtro ORDINARIO
            if applica_filtro and not str(f["TEMA"]).startswith(codici_ammessi):
                continue

            foglio = int(f["FOGLIO"])
            allegato = f["ALLEGATO"]
            mappale = f["MAPPALE"]
            gruppo_corrente = (foglio, mappale)

            # Nuovo gruppo Foglio–Mappale
            if ultimo_gruppo is None or gruppo_corrente != ultimo_gruppo:

                # Crea sempre una nuova riga per il gruppo
                sep_row = table.add_row()

                # Unisci celle
                first_cell = sep_row.cells[0]
                for c in sep_row.cells[1:]:
                    first_cell.merge(c)

                # Testo Foglio–Mappale
                testo = f"Foglio {foglio} – Mappale {mappale}"
                first_cell.text = testo

                # Formattazione testo
                for paragraph in first_cell.paragraphs:
                    run = paragraph.runs[0]
                    run.bold = True
                    run.font.name = "Calibri"
                    run.font.size = Pt(11)
                    paragraph.paragraph_format.space_before = Pt(6)

                # Rimuovi eventuale sfondo
                tcPr = first_cell._element.get_or_add_tcPr()
                for shd in tcPr.findall(".//w:shd", namespaces={"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}):
                    tcPr.remove(shd)

                # Bordi: SOLO bordo inferiore
                borders = parse_xml(r'''
                    <w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                        <w:top w:val="nil"/>
                        <w:bottom w:val="single" w:sz="6"/>
                        <w:left w:val="nil"/>
                        <w:right w:val="nil"/>
                    </w:tcBorders>
                ''')
                tcPr.append(borders)

                first_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

                # Intestazione dopo ogni gruppo
                aggiungi_intestazione()

                ultimo_gruppo = gruppo_corrente

            # Riga dati
            row = table.add_row()
            valori = [
                str(foglio),
                str(allegato),
                str(mappale),
                str(f["TEMA"]),
                str(f["ZONA"]),
                str(f["DETTAGLIO"]),
                str(f["NORME"]),
                str(f["PERCENT_V"])
            ]

            for i, cell in enumerate(row.cells):
                cell.text = valori[i]

                for paragraph in cell.paragraphs:
                    run = paragraph.runs[0]
                    run.font.name = "Calibri Light"
                    run.font.size = Pt(10)
                    paragraph.alignment = allineamenti[i]

                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        # 🔥 Elimina la riga fantasma del template (solo se è ancora vuota)
        if table.rows and all(cell.text.strip() == "" for cell in table.rows[0].cells):
            table._tbl.remove(table.rows[0]._tr)

        # Salvataggio
        doc.save(output_file_path)

        iface.messageBar().pushMessage(
            "CDU generato",
            f"Documento salvato: <a href=\"{output_file_path}\">{output_file_path}</a>",
            level=0
        )

        # Apertura automatica
        if platform.system() == "Windows":
            os.startfile(output_file_path)
        elif platform.system() == "Darwin":
            os.system(f"open {output_file_path}")
        elif platform.system() == "Linux":
            os.system(f"xdg-open {output_file_path}")

    except Exception as e:
        iface.messageBar().pushMessage("Errore", f"Errore durante la generazione del CDU: {e}", level=3)

# Esegui la funzione
export_cdu_from_destination_layer()

