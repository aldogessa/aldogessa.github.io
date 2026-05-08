from qgis.PyQt.QtWidgets import QFileDialog, QMessageBox
from qgis.core import QgsProject
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import os
import platform
from qgis.utils import iface

def export_cdu_from_destination_layer():
    try:
        # Recupera il layer su cui è installata l'azione
        layer_id = "[% @layer_id %]"
        layer = QgsProject.instance().mapLayer(layer_id)

        if not layer:
            iface.messageBar().pushMessage("Errore", "Impossibile identificare il layer dell'azione.", level=3)
            return

        # Controllo selezione
        selected = layer.selectedFeatures()
        if not selected:
            iface.messageBar().pushMessage("Avviso", "Nessun record selezionato nel layer di destinazione.", level=2)
            return

        # 🔥 Scelta ORDINARIO / COMPLETO
        msg = QMessageBox()
        msg.setWindowTitle("Tipo di certificato")
        msg.setText("Desideri il certificato ORDINARIO (filtrato) o COMPLETO (tutti i dati)?")

        btn_ordinario = msg.addButton("ORDINARIO", QMessageBox.ActionRole)
        btn_completo = msg.addButton("COMPLETO", QMessageBox.ActionRole)
        btn_annulla = msg.addButton("ANNULLA", QMessageBox.RejectRole)

        msg.exec_()

        if msg.clickedButton() == btn_annulla:
            iface.messageBar().pushMessage("Informazione", "Operazione annullata.", level=1)
            return

        applica_filtro = (msg.clickedButton() == btn_ordinario)

        # Finestra salvataggio
        output_file_path, _ = QFileDialog.getSaveFileName(
            None,
            "Scegli dove salvare il CDU",
            "",
            "Documenti Word (*.docx)"
        )
        if not output_file_path:
            iface.messageBar().pushMessage("Informazione", "Operazione annullata.", level=1)
            return

        # Template Word
        template_path = os.path.join(QgsProject.instance().homePath(), "Immagini", "00_CDU_Schema.docx")
        if not os.path.exists(template_path):
            iface.messageBar().pushMessage("Errore", f"Template non trovato: {template_path}", level=3)
            return

        doc = Document(template_path)

        # Tabella Word
        if not doc.tables:
            iface.messageBar().pushMessage("Errore", "Nessuna tabella trovata nel template!", level=3)
            return

        table = doc.tables[0]

        # Intestazioni
        headers = ["Fg.", "All.", "Map.", "Tema", "Zona", "Dettaglio", "Norme Specifiche", "Q.tà* %"]

        # Formattazione intestazione iniziale
        header_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            header_cells[i].text = header
            for paragraph in header_cells[i].paragraphs:
                run = paragraph.runs[0]
                run.bold = True
                run.font.name = "Calibri Light"
                run.font.size = Pt(10)
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D3D3D3"/>')
            header_cells[i]._element.get_or_add_tcPr().append(shading)
            header_cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        # Codici ammessi per ORDINARIO
        codici_ammessi = ("T01", "T02", "T03")

        # Ordinamento dei record selezionati
        selected_sorted = sorted(
            selected,
            key=lambda f: (
                int(f["FOGLIO"]),
                f["MAPPALE"],
                f["TEMA"],
                f["ZONA"],
                f["PERCENT_V"]
            )
        )

        ultimo_gruppo = "INIT"

        # Popolamento tabella
        for f in selected_sorted:

            # Filtro ORDINARIO
            if applica_filtro:
                if not str(f["TEMA"]).startswith(codici_ammessi):
                    continue

            foglio = int(f["FOGLIO"])
            allegato = f["ALLEGATO"]
            mappale = f["MAPPALE"]
            gruppo_corrente = (foglio, mappale)

            # Reinserisci intestazione quando cambia particella
            if ultimo_gruppo != "INIT" and gruppo_corrente != ultimo_gruppo:
                header_row = table.add_row()
                for i, header in enumerate(headers):
                    header_row.cells[i].text = header
                    for paragraph in header_row.cells[i].paragraphs:
                        run = paragraph.runs[0]
                        run.bold = True
                        run.font.name = "Calibri Light"
                        run.font.size = Pt(10)
                    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D3D3D3"/>')
                    header_row.cells[i]._element.get_or_add_tcPr().append(shading)
                    header_row.cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

            ultimo_gruppo = gruppo_corrente

            # Riga dati
            row = table.add_row()
            row.cells[0].text = str(foglio)
            row.cells[1].text = str(allegato)
            row.cells[2].text = str(mappale)
            row.cells[3].text = str(f["TEMA"])
            row.cells[4].text = str(f["ZONA"])
            row.cells[5].text = str(f["DETTAGLIO"])
            row.cells[6].text = str(f["NORME"])
            row.cells[7].text = str(f["PERCENT_V"])

            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    run = paragraph.runs[0]
                    run.font.name = "Calibri Light"
                    run.font.size = Pt(10)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        # Salvataggio
        doc.save(output_file_path)

        iface.messageBar().pushMessage(
            "CDU generato",
            f'Documento salvato: <a href="{output_file_path}">{output_file_path}</a>',
            level=0
        )

        # Apertura automatica
        if platform.system() == "Windows":
            os.startfile(output_file_path)
        elif platform.system() == "Darwin":
            os.system(f"open {output_file_path}")
        elif platform.system() == "Linux":
            os.system(f"xdg-open {output_file_path}")

    except Exception as e:
        iface.messageBar().pushMessage("Errore", f"Errore durante la generazione del CDU: {e}", level=3)

# Esegui la funzione
export_cdu_from_destination_layer()
