"""
===================================================================================
VERSIONE TESTATA SU QGIS 3.28
Questo script acquisisce i dati dalla TabellaCDU e genera il certificato di
destinazione urbanistica in formato .doc utilizzando un template preconfigurato
e accessibile a QGIS.
Script progettato per essere utilizzato come funzione definita nel layer
"TabellaCDU"
==================================================================================
"""

from qgis.PyQt.QtWidgets import QFileDialog
from qgis.core import QgsProject
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import os
import platform  # Necessario per determinare il sistema operativo
from qgis.utils import iface  # Per visualizzare i messaggi nello spazio dei messaggi QGIS

def export_selected_data_to_existing_table(template_path):
    try:
        # Finestra di dialogo per scegliere il salvataggio
        output_file_path, _ = QFileDialog.getSaveFileName(
            None,
            "Scegli dove salvare il documento",
            "",
            "Documenti Word (*.docx)"
        )

        if not output_file_path:
            iface.messageBar().pushMessage("Informazione", "Operazione annullata dall'utente.", level=1)
            return

        # Carica il modello Word
        doc = Document(template_path)

        # Recupera il layer e i dati selezionati
        layer = QgsProject.instance().mapLayersByName("TabellaCDU")[0] # Qui definire il nome del layer che contiene le intersezioni
        selected_features = layer.selectedFeatures()

        if not selected_features:
            iface.messageBar().pushMessage("Avviso", "Nessun record selezionato!", level=2)
            return

        # Ordina i dati in base alle colonne desiderate
        sorted_features = sorted(
            selected_features,
            key=lambda f: (int(f["foglio"]), f["mappale"], f["tema"], f["zona"], f["percent"])
        )

        # Trova la prima tabella nel documento
        if not doc.tables:
            iface.messageBar().pushMessage("Errore", "Nessuna tabella trovata nel documento!", level=3)
            return

        table = doc.tables[0]  # Seleziona la prima tabella

        # Imposta l'intestazione della tabella
        header_cells = table.rows[0].cells
        headers = ["Fg.", "All.", "Map.", "Tema", "Zona", "Dettaglio", "Q.tà* %"]
        for i, header in enumerate(headers):
            header_cells[i].text = header
            for paragraph in header_cells[i].paragraphs:
                run = paragraph.runs[0]
                run.bold = True  # Imposta il testo in grassetto
                run.font.name = "Calibri Light"  # Ripristina il carattere Calibri Light
                run.font.size = Pt(10)  # Dimensione del font

            # Aggiungi lo sfondo grigio alle celle dell'intestazione
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D3D3D3"/>')
            header_cells[i]._element.get_or_add_tcPr().append(shading)

            # Imposta l'allineamento verticale al centro per le celle dell'intestazione
            header_cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        # Popola la tabella con i dati selezionati
        for feature in sorted_features:
            row = table.add_row()  # Aggiungi una nuova riga
            row.cells[0].text = str(int(feature["foglio"]))  # Converti il foglio a intero
            row.cells[1].text = str(feature["allegato"])
            row.cells[2].text = str(feature["mappale"])
            row.cells[3].text = str(feature["tema"])
            row.cells[4].text = str(feature["zona"])
            row.cells[5].text = str(feature["descrizion"])
            row.cells[6].text = str(feature["percent"])

            # Formattazione delle celle
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    run = paragraph.runs[0]
                    run.font.name = "Calibri Light"
                    run.font.size = Pt(10)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER  # Allineamento verticale al centro

        # Salva il file modificato
        doc.save(output_file_path)
        iface.messageBar().pushMessage(
            "Informazione",
            f'Documento salvato con successo: <a href="{output_file_path}">{output_file_path}</a>',
            level=0
        )

        # Apertura automatica del file Word salvato
        if platform.system() == "Windows":
            os.startfile(output_file_path)
        elif platform.system() == "Darwin":  # macOS
            os.system(f"open {output_file_path}")
        elif platform.system() == "Linux":
            os.system(f"xdg-open {output_file_path}")
        else:
            iface.messageBar().pushMessage("Avviso", "Apertura automatica non supportata su questo sistema operativo.", level=2)

    except Exception as e:
        iface.messageBar().pushMessage("Errore", f"Errore durante l'esportazione: {e}", level=3)

# Specifica il percorso del modello Word
template_file = r"C:\ALDO\02_LAVORO\03_PROVE_CDU\PULA\CDU\00_CDU_Schema.docx" # Qui definire il percorso del template preconfigurato

# Esegui la funzione
export_selected_data_to_existing_table(template_file)
