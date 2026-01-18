"""
===========================================================================================
VERSIONE TESTATA IN QGIS 3.40
Questo script si usa nelle azioni layer.
Lo script è costruito su layer specifici e preventivamente configurati, dalla selezione
delle particelle catastali recupera le intersezioni già precedentemente eseguite ed
archiviate nel layer DEstinazioniUrbanistiche in relazione con una chiave esterna VIRTID.
Le righe recuperato sono utilizzate per popolare la tabella del template del CDU
preconfigurato. E' necessario omogeneizzare i campi della tabella del template con i
campi del layer di estrazione.
===========================================================================================
"""

from qgis.PyQt.QtWidgets import QFileDialog
from qgis.core import QgsProject, QgsFeature
from docx import Document
from docx.shared import Pt
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import os
import platform
from qgis.utils import iface

def export_related_data_to_existing_table():
    try:
        # Recupera il layer padre (_Particelle_) e verifica se esiste # Qui definire il layer delle particelle catastali
        parent_layer = QgsProject.instance().mapLayersByName("_Particelle_") # Qui definire il layer delle particelle catastali
        if not parent_layer or len(parent_layer) == 0:
            iface.messageBar().pushMessage("Errore", "Layer '_Particelle_' non trovato! Controlla il nome e riprova.", level=3)
            return
        parent_layer = parent_layer[0]

        # Controlla se ci sono elementi selezionati nel layer
        selected_features = parent_layer.selectedFeatures()
        if not selected_features:
            iface.messageBar().pushMessage("Avviso", "Nessun elemento selezionato nel Layer.", level=2)
            return

        # Finestra di dialogo per scegliere il salvataggio
        output_file_path, _ = QFileDialog.getSaveFileName(
            None,
            "Scegli dove salvare il documento",
            "",
            "Documenti Word (*.docx)"
        )

        if not output_file_path:
            iface.messageBar().pushMessage("Informazione", "Operazione annullata dall'utente.", level=1)
            return

        # Determina il percorso del file modello CDU_Schema
        template_path = os.path.join(QgsProject.instance().homePath(), "Immagini", "00_CDU_Schema.docx") # Qui definire il percorso del template del .doc
        if not os.path.exists(template_path):
            iface.messageBar().pushMessage("Errore", f"File modello non trovato: {template_path}", level=3)
            return

        # Carica il modello Word
        doc = Document(template_path)

        # Recupera il layer figlio (DestinazioniUrbanistiche) e verifica
        child_layer = QgsProject.instance().mapLayersByName("DestinazioniUrbanistiche") # Qui definire il layer che contiene le intersezioni preventivamente eseguite
        if not child_layer or len(child_layer) == 0:
            iface.messageBar().pushMessage("Errore", "Layer 'DestinazioniUrbanistiche' non trovato! Controlla il nome e riprova.", level=3)
            return
        child_layer = child_layer[0]

        # Trova la prima tabella nel documento
        if not doc.tables:
            iface.messageBar().pushMessage("Errore", "Nessuna tabella trovata nel documento!", level=3)
            return

        table = doc.tables[0]  # Seleziona la prima tabella

        # Imposta l'intestazione della tabella
        header_cells = table.rows[0].cells
        headers = ["Fg.", "All.", "Map.", "Tema", "Zona", "Dettaglio", "Norme Specifiche", "Q.tà* %"] # Qui definisci l'intestazione dei campi della tabella nel .doc
        for i, header in enumerate(headers):
            header_cells[i].text = header
            for paragraph in header_cells[i].paragraphs:
                run = paragraph.runs[0]
                run.bold = True
                run.font.name = "Calibri Light"
                run.font.size = Pt(10)
            shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D3D3D3"/>')
            header_cells[i]._element.get_or_add_tcPr().append(shading)
            header_cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        # Popola la tabella con i dati correlati
        for parent_feature in selected_features:
            virtid = parent_feature["VIRTID"] # Qui definisci i campi chiave esterna figlio (virtid)-padre (VIRTID)
            # Ordina i dati correlati
            related_features = sorted(
                [f for f in child_layer.getFeatures() if f["VIRTID"] == virtid],
                key=lambda feature: (
                    int(parent_feature["foglio"]),  # Layer padre: foglio
                    feature["mappale"],
                    feature["fonte"],
                    feature["zona"],
                    feature["percent"]
                )
            )

            for feature in related_features:
                row = table.add_row()
                row.cells[0].text = str(int(parent_feature["foglio"]))  # Dal layer padre
                row.cells[1].text = str(feature["allegato"])
                row.cells[2].text = str(feature["mappale"])
                row.cells[3].text = str(feature["fonte"])
                row.cells[4].text = str(feature["zona"])
                row.cells[5].text = str(feature["dettaglio"])
                row.cells[6].text = str(feature["CDU_norme"])
                row.cells[7].text = str(feature["percent"])

                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        run = paragraph.runs[0]
                        run.font.name = "Calibri Light"
                        run.font.size = Pt(10)
                    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

        # Salva il file modificato
        doc.save(output_file_path)
        iface.messageBar().pushMessage(
            "Informazione",
            f'Documento salvato con successo: <a href="{output_file_path}">{output_file_path}</a>',
            level=0
        )

        # Apertura automatica del file Word salvato
        if platform.system() == "Windows":
            os.startfile(output_file_path)
        elif platform.system() == "Darwin":  # macOS
            os.system(f"open {output_file_path}")
        elif platform.system() == "Linux":
            os.system(f"xdg-open {output_file_path}")
        else:
            iface.messageBar().pushMessage("Avviso", "Apertura automatica non supportata su questo sistema operativo.", level=2)

    except Exception as e:
        iface.messageBar().pushMessage("Errore", f"Errore durante l'esportazione: {e}", level=3)

# Esegui la funzione
export_related_data_to_existing_table()
